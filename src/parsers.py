from __future__ import annotations

import io
import re
import subprocess
import tempfile
from pathlib import Path

import pandas as pd
from docx import Document


def normalize_name(value: str) -> str:
    if value is None:
        return ""

    text = str(value).strip().lower()
    text = re.sub(r"[^a-zA-Z\s]", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def normalize_pronoun(value: str) -> str:
    if value is None:
        return ""

    text = str(value).strip().lower()

    if text in {"he", "him", "his"}:
        return "he"

    if text in {"she", "her", "hers"}:
        return "she"

    if text in {"they", "them", "their", "theirs"}:
        return "they"

    return ""


def read_class_list(uploaded_file) -> pd.DataFrame:
    file_name = uploaded_file.name.lower()

    if file_name.endswith(".csv"):
        df = pd.read_csv(uploaded_file)
    elif file_name.endswith(".xlsx"):
        df = pd.read_excel(uploaded_file)
    else:
        raise ValueError("Class list must be a CSV or Excel file.")

    required_columns = ["Name", "Report Name", "Date of Birth", "Pronoun"]

    for column in required_columns:
        if column not in df.columns:
            df[column] = ""

    df = df[required_columns].copy()

    df["Name"] = df["Name"].fillna("").astype(str).str.strip()
    df["Report Name"] = df["Report Name"].fillna("").astype(str).str.strip()
    df["Pronoun"] = df["Pronoun"].fillna("").astype(str).str.strip().str.lower()

    df["Date of Birth"] = pd.to_datetime(
        df["Date of Birth"],
        errors="coerce",
        dayfirst=False,
    )

    df["name_norm"] = df["Name"].map(normalize_name)
    df["report_name_norm"] = df["Report Name"].map(normalize_name)

    return df


def read_docx_file(uploaded_file) -> str:
    file_bytes = uploaded_file.getvalue()
    document = Document(io.BytesIO(file_bytes))

    paragraphs = [paragraph.text for paragraph in document.paragraphs]

    for table in document.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            paragraphs.append("\t".join(row_text))

    return "\n".join(paragraphs)


def _read_doc_bytes_via_textutil(file_bytes: bytes, file_name: str) -> str | None:
    with tempfile.TemporaryDirectory() as temp_dir:
        input_path = Path(temp_dir) / file_name
        output_path = Path(temp_dir) / "converted.txt"

        input_path.write_bytes(file_bytes)

        try:
            subprocess.run(
                [
                    "textutil",
                    "-convert",
                    "txt",
                    "-output",
                    str(output_path),
                    str(input_path),
                ],
                check=True,
                capture_output=True,
                text=True,
            )
        except (FileNotFoundError, subprocess.CalledProcessError):
            return None

        if not output_path.exists():
            return None

        return output_path.read_text(errors="ignore")


def _read_doc_bytes_via_libreoffice(file_bytes: bytes, file_name: str) -> str | None:
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        input_path = temp_path / file_name

        input_path.write_bytes(file_bytes)

        try:
            subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to",
                    "txt:Text",
                    "--outdir",
                    str(temp_path),
                    str(input_path),
                ],
                check=True,
                capture_output=True,
                text=True,
                timeout=60,
            )
        except (
            FileNotFoundError,
            subprocess.CalledProcessError,
            subprocess.TimeoutExpired,
        ):
            return None

        converted_files = list(temp_path.glob("*.txt"))

        if not converted_files:
            return None

        return converted_files[0].read_text(errors="ignore")


def read_doc_file(uploaded_file) -> str:
    file_bytes = uploaded_file.getvalue()
    file_name = uploaded_file.name

    text = _read_doc_bytes_via_textutil(file_bytes, file_name)

    if text:
        return text

    text = _read_doc_bytes_via_libreoffice(file_bytes, file_name)

    if text:
        return text

    raise ValueError(
        "Could not read this .doc file. Please save it as a .docx file and upload it again."
    )


def read_txt_file(uploaded_file) -> str:
    file_bytes = uploaded_file.getvalue()
    return file_bytes.decode("utf-8", errors="ignore")


def read_report_file(uploaded_file) -> str:
    file_name = uploaded_file.name.lower()

    if file_name.endswith(".docx"):
        return read_docx_file(uploaded_file)

    if file_name.endswith(".doc"):
        return read_doc_file(uploaded_file)

    if file_name.endswith(".txt"):
        return read_txt_file(uploaded_file)

    raise ValueError("Report must be a .doc, .docx, or .txt file.")
